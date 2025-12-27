from contextlib import asynccontextmanager
from pathlib import Path
from fastapi import FastAPI, Request, status
from fastapi.staticfiles import StaticFiles
from fastapi.responses import HTMLResponse, JSONResponse, FileResponse
from fastapi.exceptions import RequestValidationError
from starlette.templating import Jinja2Templates
from starlette.exceptions import HTTPException as StarletteHTTPException
from app.database import init_db
from app.routes import auth, users, admin, account
from app.routes import (
    ql_instances,
    script_configs,
    earnings,
    settlements,
    wallet,
    referrals,
    stats,
    config_envs,
)
import traceback

# 获取项目根目录（app目录的父目录）
BASE_DIR = Path(__file__).resolve().parent.parent

# 静态文件和模板目录的绝对路径
STATIC_DIR = BASE_DIR / "static"
TEMPLATES_DIR = BASE_DIR / "templates"
DATA_DIR = BASE_DIR / "data"


@asynccontextmanager
async def lifespan(app: FastAPI):
    """应用生命周期管理"""
    # 启动时执行
    init_db()
    print("数据库初始化完成")
    yield
    # 关闭时执行（如果需要）


# 创建FastAPI应用
app = FastAPI(
    title="快手账号管理平台",
    version="1.0.0",
    lifespan=lifespan
)

# 挂载静态文件（使用绝对路径）
app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")

# 挂载 data 目录（用于文件下载）
if DATA_DIR.exists():
    app.mount("/data", StaticFiles(directory=str(DATA_DIR)), name="data")

# 配置模板（使用绝对路径）
templates = Jinja2Templates(directory=str(TEMPLATES_DIR))

# 全局异常处理器
@app.exception_handler(StarletteHTTPException)
async def http_exception_handler(request: Request, exc: StarletteHTTPException):
    """处理HTTP异常，确保返回JSON格式"""
    # 如果是API请求，返回JSON
    if request.url.path.startswith("/api"):
        return JSONResponse(
            status_code=exc.status_code,
            content={"detail": exc.detail}
        )
    # 对于非API请求，重新抛出异常让FastAPI默认处理
    raise exc


@app.exception_handler(RequestValidationError)
async def validation_exception_handler(request: Request, exc: RequestValidationError):
    """处理请求验证错误"""
    if request.url.path.startswith("/api"):
        return JSONResponse(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            content={"detail": exc.errors()}
        )
    # 对于非API请求，重新抛出异常
    raise exc


@app.exception_handler(Exception)
async def general_exception_handler(request: Request, exc: Exception):
    """处理所有未捕获的异常"""
    # 记录错误详情
    print(f"未处理的异常: {exc}")
    traceback.print_exc()
    
    # 如果是API请求，返回JSON格式的错误
    if request.url.path.startswith("/api"):
        return JSONResponse(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            content={
                "detail": f"服务器内部错误: {str(exc)}"
            }
        )
    # 对于非API请求，返回通用错误页面
    from fastapi import HTTPException
    raise HTTPException(
        status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
        detail="服务器内部错误，请稍后重试"
    )


# ==================== 注册API路由 ====================
app.include_router(auth.router)
app.include_router(users.router)
app.include_router(admin.router)
app.include_router(account.router)
app.include_router(ql_instances.router)
app.include_router(script_configs.router)
app.include_router(earnings.router)
app.include_router(settlements.router)
app.include_router(wallet.router)
app.include_router(referrals.router)
app.include_router(stats.router)
app.include_router(config_envs.router)


# ==================== 页面路由 ====================

@app.get("/", response_class=HTMLResponse)
async def root(request: Request):
    """根路径重定向到登录页"""
    return templates.TemplateResponse("login.html", {"request": request})


@app.get("/login", response_class=HTMLResponse)
async def login_page(request: Request):
    """登录页面"""
    return templates.TemplateResponse("login.html", {"request": request})


@app.get("/register", response_class=HTMLResponse)
async def register_page(request: Request):
    """注册页面"""
    return templates.TemplateResponse("register.html", {"request": request})


@app.get("/dashboard", response_class=HTMLResponse)
async def dashboard_page(request: Request):
    """仪表板页面"""
    return templates.TemplateResponse("dashboard_new.html", {"request": request, "active_page": "dashboard"})


@app.get("/ks-accounts", response_class=HTMLResponse)
async def ks_accounts_page(request: Request):
    """快手账号管理页面"""
    return templates.TemplateResponse("ks_accounts.html", {"request": request, "active_page": "ks_accounts"})


@app.get("/config-envs", response_class=HTMLResponse)
async def config_envs_page(request: Request):
    """配置环境管理页面"""
    return templates.TemplateResponse(
        "config_envs.html",
        {"request": request, "active_page": "config_envs"},
    )


@app.get("/earnings", response_class=HTMLResponse)
async def earnings_page(request: Request):
    """收益统计页面"""
    return templates.TemplateResponse("earnings.html", {"request": request, "active_page": "earnings"})


@app.get("/settlements", response_class=HTMLResponse)
async def settlements_page(request: Request):
    """结算管理页面"""
    return templates.TemplateResponse("settlements.html", {"request": request, "active_page": "settlements"})


@app.get("/wallet", response_class=HTMLResponse)
async def wallet_page(request: Request):
    """我的钱包页面"""
    return templates.TemplateResponse("wallet.html", {"request": request, "active_page": "wallet"})


@app.get("/referral", response_class=HTMLResponse)
async def referral_page(request: Request):
    """推广中心页面"""
    return templates.TemplateResponse("referral.html", {"request": request, "active_page": "referral"})

@app.get("/account", response_class=HTMLResponse)
async def account_page(request: Request):
    """个人账户页面"""
    return templates.TemplateResponse("account.html", {"request": request, "active_page": "account"})


# 避免 favicon 404 打出 500 日志
@app.get("/favicon.ico")
async def favicon():
    from fastapi import Response
    return Response(status_code=204)


# ==================== 文档和下载 API ====================

@app.get("/api/guide/content")
async def get_guide_content():
    """获取新手搭建说明文档内容"""
    docx_path = DATA_DIR / "describe" / "新手搭建说明.docx"

    if not docx_path.exists():
        return JSONResponse(
            status_code=404,
            content={"detail": "说明文档不存在"}
        )

    try:
        from docx import Document
        doc = Document(str(docx_path))

        content_parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # 根据段落样式判断是否为标题
                if para.style.name.startswith('Heading'):
                    level = para.style.name[-1] if para.style.name[-1].isdigit() else '2'
                    content_parts.append(f"<h{level}>{text}</h{level}>")
                else:
                    content_parts.append(f"<p>{text}</p>")

        # 处理表格
        for table in doc.tables:
            table_html = "<table class='doc-table'>"
            for row in table.rows:
                table_html += "<tr>"
                for cell in row.cells:
                    table_html += f"<td>{cell.text}</td>"
                table_html += "</tr>"
            table_html += "</table>"
            content_parts.append(table_html)

        return {"content": "\n".join(content_parts)}
    except ImportError:
        return JSONResponse(
            status_code=500,
            content={"detail": "服务器缺少 python-docx 库，请安装: pip install python-docx"}
        )
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"detail": f"读取文档失败: {str(e)}"}
        )


# ==================== 管理员页面路由 ====================

@app.get("/admin/users", response_class=HTMLResponse)
async def admin_users_page(request: Request):
    """用户管理页面（管理员）"""
    return templates.TemplateResponse("admin_users.html", {"request": request, "active_page": "users"})


@app.get("/admin/ql-instances", response_class=HTMLResponse)
async def admin_ql_instances_page(request: Request):
    """青龙实例管理页面（管理员）"""
    return templates.TemplateResponse("admin_ql_instances.html", {"request": request, "active_page": "ql_instances"})


@app.get("/admin/referrals", response_class=HTMLResponse)
async def admin_referrals_page(request: Request):
    """推广关系管理页面（管理员）"""
    return templates.TemplateResponse("admin_referrals.html", {"request": request, "active_page": "referrals"})


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=1212)
